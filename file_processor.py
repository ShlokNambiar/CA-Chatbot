import os
import pandas as pd
import docx
import PyPDF2
import pdfplumber
from typing import Dict, Any, List, Optional
import io
import json

class FileProcessor:
    """Service for processing uploaded files (PDF, Word, Excel, CSV)"""
    
    def __init__(self):
        """Initialize file processor"""
        self.supported_extensions = {
            '.pdf': self.process_pdf,
            '.docx': self.process_word,
            '.doc': self.process_word,
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.csv': self.process_csv
        }
    
    def process_file(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract content
        
        Args:
            file_path: Path to the uploaded file
            file_name: Original name of the file
            
        Returns:
            Dictionary with processed content and metadata
        """
        try:
            # Get file extension
            _, ext = os.path.splitext(file_name.lower())
            
            if ext not in self.supported_extensions:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {ext}',
                    'supported_types': list(self.supported_extensions.keys())
                }
            
            # Process file based on type
            processor = self.supported_extensions[ext]
            result = processor(file_path, file_name)
            
            return {
                'success': True,
                'file_name': file_name,
                'file_type': ext,
                'content': result.get('content', ''),
                'metadata': result.get('metadata', {}),
                'summary': self._create_summary(result, ext)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_name': file_name
            }
    
    def process_pdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process PDF file"""
        content = ""
        metadata = {}
        
        try:
            # Try pdfplumber first (better for tables and complex layouts)
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
                
                # Extract tables if any
                tables = []
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                
                if tables:
                    metadata['tables_found'] = len(tables)
                    content += "\n\nTABLES EXTRACTED:\n"
                    for i, table in enumerate(tables[:3], 1):  # Limit to first 3 tables
                        content += f"\nTable {i}:\n"
                        for row in table[:5]:  # Limit to first 5 rows
                            content += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
        
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            except Exception as e2:
                raise Exception(f"Failed to process PDF: {e2}")
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def process_word(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            doc = docx.Document(file_path)
            content = ""
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract tables
            if doc.tables:
                content += "\n\nTABLES:\n"
                for i, table in enumerate(doc.tables, 1):
                    content += f"\nTable {i}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        content += row_text + "\n"
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            raise Exception(f"Failed to process Word document: {e}")
    
    def process_excel(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = ""
            metadata = {
                'sheets': excel_file.sheet_names,
                'total_sheets': len(excel_file.sheet_names)
            }
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                content += f"\n\nSHEET: {sheet_name}\n"
                content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                content += f"Columns: {', '.join(df.columns.tolist())}\n"
                
                # Add sample data (first 5 rows)
                content += "\nSample Data:\n"
                content += df.head().to_string(index=False) + "\n"
                
                # Add basic statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    content += "\nNumeric Summary:\n"
                    content += df[numeric_cols].describe().to_string() + "\n"
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            raise Exception(f"Failed to process Excel file: {e}")
    
    def process_csv(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path)
            content = ""
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist()
            }
            
            content += f"CSV FILE ANALYSIS\n"
            content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            
            # Sample data
            content += "Sample Data (First 5 rows):\n"
            content += df.head().to_string(index=False) + "\n\n"
            
            # Data types
            content += "Data Types:\n"
            for col, dtype in df.dtypes.items():
                content += f"- {col}: {dtype}\n"
            
            # Basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += "\nNumeric Summary:\n"
                content += df[numeric_cols].describe().to_string() + "\n"
            
            # Missing values
            missing = df.isnull().sum()
            if missing.sum() > 0:
                content += "\nMissing Values:\n"
                for col, count in missing.items():
                    if count > 0:
                        content += f"- {col}: {count} missing\n"
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            raise Exception(f"Failed to process CSV file: {e}")
    
    def _create_summary(self, result: Dict[str, Any], file_type: str) -> str:
        """Create a brief summary of the processed file"""
        content = result.get('content', '')
        metadata = result.get('metadata', {})
        
        if file_type == '.pdf':
            return f"PDF with {metadata.get('pages', 0)} pages. {len(content)} characters extracted."
        elif file_type in ['.docx', '.doc']:
            return f"Word document with {metadata.get('paragraphs', 0)} paragraphs and {metadata.get('tables', 0)} tables."
        elif file_type in ['.xlsx', '.xls']:
            return f"Excel file with {metadata.get('total_sheets', 0)} sheets: {', '.join(metadata.get('sheets', []))}"
        elif file_type == '.csv':
            return f"CSV file with {metadata.get('rows', 0)} rows and {metadata.get('columns', 0)} columns."
        
        return f"File processed successfully. {len(content)} characters extracted."
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.supported_extensions.keys())
